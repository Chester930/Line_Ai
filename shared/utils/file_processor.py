import os
import tempfile
import asyncio
from typing import List, Dict, Optional, Tuple, Union, Any, BinaryIO
import docx
import pandas as pd
import PyPDF2
from PIL import Image
import speech_recognition as sr
import io
import logging
import hashlib
from datetime import datetime
from langdetect import detect
from jieba import analyse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from textblob import TextBlob
from whoosh.fields import Schema, TEXT, ID, DATETIME, KEYWORD
from whoosh.analysis import StemmingAnalyzer
from whoosh import index
from whoosh.qparser import QueryParser
import spacy
from shared.database.crud import DocumentCRUD
from shared.database.models import Document
from shared.ai.media_agent import MediaAgent
import shutil
import mimetypes
from pathlib import Path

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self, db=None, upload_dir: str = 'uploads'):
        self.db = db
        self.temp_dir = tempfile.mkdtemp()
        self.media_agent = MediaAgent()
        self.MAX_FILE_SIZE = 300 * 1024 * 1024  # LINE限制：300MB
        
        # 初始化 spaCy
        try:
            self.nlp = spacy.load('zh_core_web_sm')
        except:
            self.nlp = None
            logger.warning("無法載入 spaCy 中文模型")
        
        # 初始化全文搜索索引
        self.index_dir = os.path.join('data', 'search_index')
        os.makedirs(self.index_dir, exist_ok=True)
        
        self.schema = Schema(
            id=ID(stored=True),
            title=TEXT(stored=True),
            content=TEXT(analyzer=StemmingAnalyzer()),
            file_type=TEXT(stored=True),
            created_at=DATETIME(stored=True),
            keywords=KEYWORD(stored=True, commas=True)
        )
        
        # 創建或打開搜索索引
        if not index.exists_in(self.index_dir):
            self.ix = index.create_in(self.index_dir, self.schema)
        else:
            self.ix = index.open_dir(self.index_dir)
        
        # 初始化 TF-IDF 向量化器
        self.vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english'
        )
        
        # 初始化結巴分詞的 TF-IDF 關鍵詞提取器
        self.tfidf = analyse.extract_tags
        
        self.upload_dir = upload_dir
        os.makedirs(self.upload_dir, exist_ok=True)
        
        # 支持的文件類型
        self.supported_types = {
            'text/plain': self._process_text,
            'application/pdf': self._process_text,  # 暫時作為文本處理
            'application/msword': self._process_text,  # 暫時作為文本處理
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_text
        }
        
        # 定義不同文件類型的處理選項
        self.processing_options = {
            'text/plain': {'chunk_size': 1000, 'overlap': 200, 'min_length': 50},
            'application/pdf': {'chunk_size': 1500, 'overlap': 300, 'min_length': 100},
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': {'chunk_size': 1200, 'overlap': 250, 'min_length': 75},
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': {'chunk_size': 800, 'overlap': 150, 'min_length': 40},
            'application/json': {'chunk_size': 500, 'overlap': 100, 'min_length': 30},
            'text/markdown': {'chunk_size': 1000, 'overlap': 200, 'min_length': 50},
            'text/html': {'chunk_size': 1000, 'overlap': 200, 'min_length': 50}
        }

    def _clean_and_format_content(self, content: str) -> str:
        """清理和格式化內容"""
        import re
        
        # 基本清理
        content = content.strip()
        
        # 移除多餘的空白字符
        content = ' '.join(content.split())
        
        # 標準化換行符
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除重複的換行
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # 確保句子之間有適當的空格
        content = content.replace('.', '. ').replace('。', '。 ')
        content = re.sub(r'\s{2,}', ' ', content)
        
        # 移除特殊字符和控制字符
        content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', content)
        
        # 修復常見的格式問題（修正括號問題）
        content = re.sub(r'([。！？!?])([^"\'"\'])', r'\1\n\2', content)  # 在句號後添加換行
        content = re.sub(r'([。！？!?]["\'"\'"])([^，。！？!?])', r'\1\n\2', content)  # 在引號後添加換行
        
        # 移除空行開頭的空格
        content = re.sub(r'\n\s+', '\n', content)
        
        # 確保段落之間有適當的間距
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        content = '\n\n'.join(paragraphs)
        
        return content

    def _structure_content(self, content: str, file_type: str) -> Dict[str, Any]:
        """將內容結構化，便於後續處理和存儲"""
        options = self.processing_options.get(file_type, {
            'chunk_size': 1000,
            'overlap': 200,
            'min_length': 50
        })
        
        # 分割內容為段落
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # 創建文檔塊
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) <= options['chunk_size']:
                current_chunk += (paragraph + "\n\n")
            else:
                if len(current_chunk) >= options['min_length']:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk and len(current_chunk) >= options['min_length']:
            chunks.append(current_chunk.strip())
        
        # 提取關鍵信息
        import re
        title_candidates = re.findall(r'^[第一二三四五六七八九十\d]+[章節]\s*(.+)$', content, re.MULTILINE)
        headings = re.findall(r'^#+\s*(.+)$', content, re.MULTILINE)
        
        return {
            'chunks': chunks,
            'metadata': {
                'total_chunks': len(chunks),
                'average_chunk_size': sum(len(c) for c in chunks) / len(chunks) if chunks else 0,
                'titles': title_candidates[:5],  # 取前5個可能的標題
                'headings': headings[:10],  # 取前10個標題
                'total_length': len(content),
                'processing_options': options
            }
        }

    async def process_file_with_timeout(
        self,
        file: Union[str, bytes, io.BytesIO],
        file_type: str,
        timeout: int = 25,
        save_to_db: bool = False
    ) -> Dict:
        """處理檔案（帶有超時限制）
        
        Args:
            file: 檔案內容（可以是路徑、bytes或BytesIO）
            file_type: 檔案MIME類型
            timeout: 處理超時時間（秒）
            save_to_db: 是否保存到資料庫
            
        Returns:
            Dict: 處理結果
        """
        try:
            # 檢查檔案大小
            if isinstance(file, (bytes, io.BytesIO)):
                file_size = len(file.getvalue()) if isinstance(file, io.BytesIO) else len(file)
                if not self.check_file_size(file_size):
                    raise ValueError(f"檔案大小超過限制（最大 {self.MAX_FILE_SIZE/1024/1024}MB）")
            
            # 使用 asyncio 處理超時
            result = await asyncio.wait_for(
                self._process_file_async(file, file_type, save_to_db),
                timeout=timeout
            )
            
            return result
            
        except asyncio.TimeoutError:
            logger.error("檔案處理超時")
            return {
                'success': False,
                'error': "處理超時，請稍後再試"
            }
        except Exception as e:
            logger.error(f"處理檔案時發生錯誤：{str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def _process_file_async(
        self,
        file: Union[str, bytes, io.BytesIO],
        file_type: str,
        save_to_db: bool = False
    ) -> Dict:
        try:
            # 讀取文件內容
            if isinstance(file, (str, bytes)):
                file_content = file if isinstance(file, bytes) else file.encode('utf-8')
            else:
                file_content = file.read()

            # 處理文件內容
            processor = self.supported_types.get(file_type, self._process_text)
            content = processor(io.BytesIO(file_content))
            
            # 結構化內容
            structured_content = self._structure_content(content, file_type)
            
            # 準備結果
            result = {
                'success': True,
                'content': content,
                'processed_content': '\n'.join(structured_content['chunks']),
                'metadata': structured_content['metadata'],
                'file_type': file_type,
                'file_size': len(file_content)
            }

            if save_to_db and self.db:
                # 計算內容雜湊值
                content_hash = hashlib.sha256(result['processed_content'].encode('utf-8')).hexdigest()
                
                # 檢查是否存在相同內容的文件
                existing_doc = self.db.query(Document).filter(
                    Document.content_hash == content_hash
                ).first()
                
                if existing_doc:
                    # 檢測變化
                    changes = self._detect_content_changes(
                        existing_doc.processed_content or existing_doc.content,
                        result['processed_content']
                    )
                    
                    # 如果內容有顯著變化，創建新版本
                    if changes['diff_ratio'] < 0.95:  # 如果差異超過5%
                        # 保存新版本
                        file_path = self._save_file(file_content, file_type)
                        doc = DocumentCRUD().create_document(
                            title=f"{existing_doc.title}_v{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                            content=result['content'],
                            processed_content=result['processed_content'],
                            file_path=file_path,
                            file_type=file_type,
                            file_size=result['file_size'],
                            doc_metadata={
                                **result['metadata'],
                                'version_info': {
                                    'original_id': existing_doc.id,
                                    'changes': changes
                                }
                            },
                            content_hash=content_hash
                        )
                        result['document_id'] = doc.id
                        result['version_info'] = {
                            'is_new_version': True,
                            'original_id': existing_doc.id,
                            'changes': changes
                        }
                    else:
                        # 內容相似度高，不創建新版本
                        result['document_id'] = existing_doc.id
                        result['version_info'] = {
                            'is_new_version': False,
                            'message': '文件內容相似度高，未創建新版本'
                        }
                else:
                    # 創建新文件
                    file_path = self._save_file(file_content, file_type)
                    doc = DocumentCRUD().create_document(
                        title=structured_content['metadata']['titles'][0] if structured_content['metadata']['titles'] else f"uploaded_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                        content=result['content'],
                        processed_content=result['processed_content'],
                        file_path=file_path,
                        file_type=file_type,
                        file_size=result['file_size'],
                        doc_metadata=result['metadata'],
                        content_hash=content_hash
                    )
                    result['document_id'] = doc.id
            
            # 查找相似文件
            similar_docs = self.find_similar_documents(result['processed_content'])
            if similar_docs:
                result['similar_documents'] = similar_docs
            
            return result
            
        except Exception as e:
            logger.error(f"處理文件時發生錯誤: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def _extract_content_async(self, file: io.BytesIO, file_type: str) -> Dict:
        """非同步提取檔案內容"""
        try:
            # 圖片檔案
            if file_type.startswith('image/'):
                image = Image.open(file)
                result = await self.media_agent.process_media_async(image, 'image')
                if result['success']:
                    return {'type': 'text', 'text': result['text']}
                else:
                    raise ValueError(result['error'])
            
            # 音訊檔案
            elif file_type.startswith('audio/'):
                result = await self.media_agent.process_media_async(file.getvalue(), 'audio')
                if result['success']:
                    return {'type': 'text', 'text': result['text']}
                else:
                    raise ValueError(result['error'])
            
            # 其他檔案類型處理（同步處理）
            return await asyncio.to_thread(self._extract_content_sync, file, file_type)
            
        except Exception as e:
            raise ValueError(f"提取內容時發生錯誤：{str(e)}")
    
    def _extract_content_sync(self, file: io.BytesIO, file_type: str) -> Dict:
        """同步提取檔案內容"""
        try:
            # 文本檔案
            if file_type == "text/plain":
                text = file.getvalue().decode('utf-8')
                return {'type': 'text', 'text': text}
            
            # PDF 檔案
            elif file_type == "application/pdf":
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return {'type': 'pdf', 'text': text}
            
            # Word 檔案
            elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                doc = docx.Document(file)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return {'type': 'word', 'text': text}
            
            # Excel 檔案
            elif file_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                df = pd.read_excel(file)
                text = "Excel 檔案內容：\n"
                text += f"表格大小：{df.shape[0]}行 x {df.shape[1]}列\n"
                text += f"欄位名稱：{', '.join(df.columns)}\n"
                text += "\n前5行資料：\n"
                text += df.head().to_string()
                return {'type': 'excel', 'text': text}
            
            else:
                raise ValueError(f"不支援的檔案類型：{file_type}")
            
        except Exception as e:
            raise ValueError(f"處理檔案內容時發生錯誤：{str(e)}")
    
    def check_file_size(self, size: int) -> bool:
        """檢查檔案大小是否符合限制"""
        return size <= self.MAX_FILE_SIZE
    
    def _save_file(self, file: io.BytesIO, file_type: str) -> str:
        """保存檔案到永久存儲"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        extension = self._get_file_extension(file_type)
        filename = f"{timestamp}{extension}"
        file_path = os.path.join(self.upload_dir, filename)
        
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        with open(file_path, 'wb') as f:
            f.write(file.getvalue())
            
        return file_path
    
    def _get_file_extension(self, file_type: str) -> str:
        """根據MIME類型獲取檔案副檔名"""
        extensions = {
            'text/plain': '.txt',
            'application/pdf': '.pdf',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/vnd.ms-excel': '.xls',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'audio/wav': '.wav',
            'audio/mp3': '.mp3'
        }
        return extensions.get(file_type, '')
    
    def __del__(self):
        """清理臨時檔案"""
        import shutil
        try:
            shutil.rmtree(self.temp_dir)
        except:
            pass

    def _process_text(self, file: BinaryIO) -> str:
        """處理文本文件"""
        try:
            content = file.read()
            # 嘗試不同的編碼
            encodings = ['utf-8', 'big5', 'gb18030']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
                    
            raise ValueError("Unable to decode file with supported encodings")
            
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise
    
    def _process_pdf(self, file: io.BytesIO) -> str:
        """處理 PDF 文件"""
        pdf = PyPDF2.PdfReader(file)
        text = []
        for page in pdf.pages:
            text.append(page.extract_text())
        return '\n'.join(text)
    
    def _process_docx(self, file: io.BytesIO) -> str:
        """處理 Word 文件"""
        doc = docx.Document(file)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def _process_excel(self, file: io.BytesIO) -> str:
        """處理 Excel 文件"""
        df = pd.read_excel(file)
        # 將 DataFrame 轉換為易讀的文本格式
        return df.to_string(index=False)

    def _analyze_content(self, content: str) -> Dict[str, Any]:
        """分析內容，提取關鍵信息"""
        try:
            # 檢測語言
            lang = detect(content)
            
            # 使用 jieba 提取關鍵詞
            keywords = self.tfidf(content, topK=10, withWeight=True)
            keywords = [(word, float(weight)) for word, weight in keywords]
            
            # 使用 TextBlob 進行情感分析和摘要
            blob = TextBlob(content)
            sentiment = blob.sentiment.polarity
            
            # 簡單的摘要生成（取前三句話）
            sentences = blob.sentences
            summary = " ".join(str(sentence) for sentence in sentences[:3])
            
            # 計算文本統計信息
            stats = {
                'total_chars': len(content),
                'total_words': len(content.split()),
                'total_sentences': len(sentences),
                'avg_sentence_length': len(content.split()) / len(sentences) if sentences else 0
            }
            
            return {
                'language': lang,
                'keywords': keywords,
                'sentiment': sentiment,
                'summary': summary,
                'stats': stats
            }
        except Exception as e:
            logger.error(f"分析內容時發生錯誤：{str(e)}")
            return {
                'language': 'unknown',
                'keywords': [],
                'sentiment': 0.0,
                'summary': '',
                'stats': {}
            }

    def _process_json(self, file: io.BytesIO) -> str:
        """處理 JSON 文件"""
        import json
        content = file.read().decode('utf-8')
        # 格式化 JSON 以便閱讀
        parsed = json.loads(content)
        return json.dumps(parsed, ensure_ascii=False, indent=2)

    def _process_markdown(self, file: io.BytesIO) -> str:
        """處理 Markdown 文件"""
        import markdown
        content = file.read().decode('utf-8')
        # 將 Markdown 轉換為純文本
        html = markdown.markdown(content)
        # 移除 HTML 標籤
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()

    def _process_html(self, file: io.BytesIO) -> str:
        """處理 HTML 文件"""
        from bs4 import BeautifulSoup
        content = file.read().decode('utf-8')
        soup = BeautifulSoup(content, 'html.parser')
        # 移除腳本和樣式
        for script in soup(["script", "style"]):
            script.decompose()
        return soup.get_text()

    def calculate_similarity(self, text1: str, text2: str) -> float:
        """計算兩段文本的相似度"""
        try:
            # 將文本轉換為 TF-IDF 向量
            tfidf_matrix = self.vectorizer.fit_transform([text1, text2])
            # 計算餘弦相似度
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return float(similarity)
        except Exception as e:
            logger.warning(f"計算相似度時發生錯誤：{str(e)}")
            return 0.0

    def find_similar_documents(self, content: str, threshold: float = 0.7) -> List[Dict]:
        """查找相似的文件"""
        try:
            if not self.db:
                return []
            
            similar_docs = []
            # 獲取所有文件
            documents = self.db.query(Document).all()
            
            for doc in documents:
                similarity = self.calculate_similarity(content, doc.processed_content or doc.content)
                if similarity >= threshold:
                    similar_docs.append({
                        'document_id': doc.id,
                        'title': doc.title,
                        'similarity': similarity
                    })
            
            # 按相似度排序
            similar_docs.sort(key=lambda x: x['similarity'], reverse=True)
            return similar_docs
            
        except Exception as e:
            logger.error(f"查找相似文件時發生錯誤：{str(e)}")
            return []

    def _calculate_content_hash(self, content: str) -> str:
        """計算內容的雜湊值"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()

    def _detect_content_changes(self, old_content: str, new_content: str) -> Dict[str, Any]:
        """檢測內容變化"""
        import difflib
        
        # 使用 difflib 計算差異
        differ = difflib.Differ()
        diff = list(differ.compare(old_content.splitlines(), new_content.splitlines()))
        
        # 統計變化
        added = len([line for line in diff if line.startswith('+ ')])
        removed = len([line for line in diff if line.startswith('- ')])
        changed = len([line for line in diff if line.startswith('? ')])
        
        return {
            'added_lines': added,
            'removed_lines': removed,
            'changed_lines': changed,
            'total_changes': added + removed + changed,
            'diff_ratio': difflib.SequenceMatcher(None, old_content, new_content).ratio()
        }

    def search_documents(self, query: str, field: str = 'content', limit: int = 10) -> List[Dict]:
        """全文搜索文件"""
        try:
            with self.ix.searcher() as searcher:
                query_parser = QueryParser(field, schema=self.schema)
                query_obj = query_parser.parse(query)
                results = searcher.search(query_obj, limit=limit)
                
                return [{
                    'id': result['id'],
                    'title': result['title'],
                    'file_type': result['file_type'],
                    'created_at': result['created_at'],
                    'keywords': result['keywords'].split(',') if result['keywords'] else [],
                    'score': result.score
                } for result in results]
        except Exception as e:
            logger.error(f"搜索文件時發生錯誤：{str(e)}")
            return []

    def update_search_index(self, doc: Document):
        """更新搜索索引"""
        try:
            writer = self.ix.writer()
            
            # 提取關鍵詞
            analysis_result = self._analyze_content(doc.processed_content or doc.content)
            keywords = ','.join(
                [k for k, _ in analysis_result.get('keywords', {}).get('chinese', [])] +
                list(analysis_result.get('keywords', {}).get('auto', []))
            )
            
            # 更新索引
            writer.update_document(
                id=str(doc.id),
                title=doc.title,
                content=doc.processed_content or doc.content,
                file_type=doc.file_type,
                created_at=doc.created_at,
                keywords=keywords
            )
            
            writer.commit()
            
        except Exception as e:
            logger.error(f"更新搜索索引時發生錯誤：{str(e)}")